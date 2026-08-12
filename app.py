# ==================== 导入库 ====================
import streamlit as st
import polars as pl
import matplotlib.pyplot as plt
import numpy as np
import io
from scipy.stats import chi2_contingency, ttest_ind
from statsmodels.stats.power import TTestIndPower, NormalIndPower
from statsmodels.stats.proportion import proportion_effectsize
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO

# ==================== 全局设置 ====================
# 字体设置
try:
    plt.rcParams['font.sans-serif'] = ['WenQuanYi Zen Hei', 'SimHei', 'Arial Unicode MS']
except Exception:
    pass
plt.rcParams['axes.unicode_minus'] = False

st.set_page_config(page_title='AB测试报告', layout='wide')
st.title('🚀 自动化AB测试报告生成器')
uploaded_file = st.file_uploader('上传你的CSV数据文件', type='csv')

# ==================== 文件上传后主流程 ====================
if uploaded_file is not None:
    # ---------- 1. 数据加载 ----------
    df = pl.read_csv(io.BytesIO(uploaded_file.read()))
    st.success(f'✅ 加载成功！共 {df.height} 行')

    with st.expander('查看原始数据'):
        st.dataframe(df.head(10))

    # ---------- 2. 列选择配置 ----------
    with st.expander('🎯 请指定数据列（告诉程序哪一列是什么）', expanded=True):
        cols = df.columns
        group_col = st.selectbox(
            "请选择【分组列】（例如：A组/B组）",
            options=cols,
            index=None,
            placeholder='请选择一列...'
        )
        value_col = st.selectbox(
            "请选择【数值列】（例如：时长、金额）",
            options=cols,
            index=None,
            placeholder='请选择一列...'
        )
        conv_col = st.selectbox(
            "请选择【转化列】（0/1 二分类）",
            options=cols,
            index=None,
            placeholder='请选择一列'
        )
        st.info(f"✅ 当前配置：分组={group_col}，数值={value_col}，转化={conv_col}")

    # ---------- 3. 核心分析（仅当三列都选好后执行）----------
    if group_col is not None and value_col is not None and conv_col is not None:

        # ------- 3.1 数据清洗与验证 -------
        conv_clean = df[conv_col].drop_nulls()
        if conv_clean.n_unique() > 2 or not all((conv_clean == 0) | (conv_clean == 1)):
            st.error("❌ 转化列必须只包含 0 和 1（或空值），请检查数据后重新上传。")
            st.stop()
        df = df.with_columns(pl.col(conv_col).fill_null(0).alias(conv_col))

        # ------- 3.2 基本汇总与分组检查 -------
        summary = df.group_by(group_col).agg([
            pl.len().alias('样本量'),
            (pl.col(conv_col).mean() * 100).alias('转化率%'),
            pl.col(value_col).mean().alias('平均值')
        ])
        unique_groups = df[group_col].unique().to_list()
        if len(unique_groups) != 2:
            st.error(f"❌ 分组列 '{group_col}' 包含 {len(unique_groups)} 个不同值，需要恰好两组。")
            st.stop()
        group_a_name = unique_groups[0]
        group_b_name = unique_groups[1]

        # ------- 3.3 数值列 T 检验与效应量 -------
        val_a = df.filter(pl.col(group_col) == group_a_name)[value_col].drop_nulls().to_numpy()
        val_b = df.filter(pl.col(group_col) == group_b_name)[value_col].drop_nulls().to_numpy()
        t_stat, p_value = ttest_ind(val_a, val_b, equal_var=False)

        mean_a = np.mean(val_a)
        mean_b = np.mean(val_b)
        n1, n2 = len(val_a), len(val_b)
        var_a = np.var(val_a, ddof=1)
        var_b = np.var(val_b, ddof=1)
        pooled_std = np.sqrt(((n1 - 1) * var_a + (n2 - 1) * var_b) / (n1 + n2 - 2)) if (n1 + n2 - 2) > 0 else 0
        cohens_d = abs(mean_b - mean_a) / pooled_std if pooled_std != 0 else 0

        # ------- 3.4 展示基本指标与 T 检验结果 -------
        col1, col2 = st.columns(2)
        with col1:
            st.subheader("📊 分组指标")
            st.dataframe(summary.to_pandas())
        with col2:
            st.subheader("📈 检验结论")
            st.metric('P值', f'{p_value:.4f}')
            if p_value < 0.05:
                st.success(f"✅ 结论：{group_a_name}组 与 {group_b_name}组 差异 **显著**")
            else:
                st.info(f"ℹ️ 结论：{group_a_name}组 与 {group_b_name}组 差异 **不显著**")

        # ------- 3.5 转化率卡方检验 -------
        st.subheader("📊 转化率深度拆解")
        cross_tab = df.group_by(group_col).agg([
            pl.len().alias("总人数"),
            pl.col(conv_col).sum().alias("转化人数"),
        ])
        cross_tab = cross_tab.with_columns(
            ((pl.col("转化人数") / pl.col("总人数")) * 100).alias("转化率(%)")
        )

        converted = cross_tab["转化人数"].to_list()
        not_converted = (cross_tab["总人数"] - cross_tab["转化人数"]).to_list()
        converted = [int(x) for x in converted]
        not_converted = [int(x) for x in not_converted]
        contingency_table = [converted, not_converted]

        if any(sum(row) == 0 for row in contingency_table) or \
           all(converted[i] == 0 and not_converted[i] == 0 for i in range(len(converted))):
            st.warning("⚠️ 转化数据中存在全零行/列，卡方检验可能不可靠。")
            p_value_conv = None
        else:
            try:
                chi2, p_value_conv, dof, expected = chi2_contingency(contingency_table)
            except ValueError as e:
                st.error(f"卡方检验失败：{e}")
                st.stop()

        rate_a = cross_tab.filter(pl.col(group_col) == group_a_name)["转化率(%)"].to_list()[0]
        rate_b = cross_tab.filter(pl.col(group_col) == group_b_name)["转化率(%)"].to_list()[0]

        # 对照组识别
        if "对照" in group_a_name or "Control" in group_a_name:
            rate_control, rate_treatment = rate_a, rate_b
        elif "对照" in group_b_name or "Control" in group_b_name:
            rate_control, rate_treatment = rate_b, rate_a
        else:
            if rate_a < rate_b:
                rate_control, rate_treatment = rate_a, rate_b
            else:
                rate_control, rate_treatment = rate_b, rate_a

        if rate_control > 0:
            lift = ((rate_treatment - rate_control) / rate_control) * 100
        else:
            lift = None

        st.dataframe(cross_tab.to_pandas())
        col_conv1, col_conv2 = st.columns(2)
        with col_conv1:
            if p_value_conv is not None:
                st.metric("卡方检验 P值", f"{p_value_conv:.4f}")
                if p_value_conv < 0.05:
                    st.success("✅ 转化率差异显著")
                else:
                    st.info("ℹ️ 转化率差异不显著")
            else:
                st.info("未计算P值（数据异常）")
        with col_conv2:
            lift_text = f"{lift:.1f}%" if lift is not None else "N/A"
            delta_text = f"{rate_treatment - rate_control:.1f} 个百分点"
            st.metric("相对提升幅度", lift_text, delta=delta_text)
            st.caption(f"{group_a_name} 转化率: {rate_a:.1f}% → {group_b_name} 转化率: {rate_b:.1f}%")

        #----------------------------
        with st.expander('SRM 检验(样本比率匹配检验)'):
            st.markdown("""
            ***SRM 检验***: 验证实际样本分配是否复合预期分流比例。
            如果P值 < 0.05, 说明实验可能存在分流异常， 后续所有结论需谨慎对待
            """)
            n_a = len(val_a)
            n_b = len(val_b)
            total_n = n_a + n_b
            col_srm1, col_srm2 = st.columns(2)
            with col_srm1:
                expected_ratio_a = st.slider(
                    "A组期望分流比例(%)",
                    min_value=10, max_value=90, value=50, step=5
                ) / 100.0

                expected_ratio_b = 1 - expected_ratio_a
                st.caption(f'B 组比例自动为{expected_ratio_b*100:.0f}%')

            expected_a = total_n * expected_ratio_a
            expected_b = total_n * expected_ratio_b
            with col_srm2:
                st.metric('实际样本量', f"A组:{n_a}人, B组:{n_b}人")
                st.metric("期望样本量",
                          f"A组{int(expected_a)}人，"
                          f"B组{int(expected_b)}人，")

            observed = [n_a, n_b]
            expected = [expected_a, expected_b]

            chi2_srm, p_value_srm = chi2_contingency([observed, expected])[:2]

            st.divider()
            col_res1, col_res2 = st.columns(2)
            with col_res1:
                st.metric('SRM 卡方值', f"{chi2_srm:.4f}")
                st.metric('SRM P值', f"{p_value_srm:.4f}")
            with col_res2:
                if p_value_srm >= 0.05:
                    st.success('✅ SRM 检验通过(p>=0.05)')
                    st.info('样本分流复合预期, 数据质量可靠, 可继续分析')
                else:
                    st.error("❌ SRM 检验未通过(p<0.05)")
                    st.warning("""
                    ⚠️ 实际样本分配偏离显著预期，可能存在以下问题:
                    - 分流系统配置错误(如A/B组流量比例设置不一致)
                    - 数据手机过程中的系统偏差(如某组用户被大量过滤)
                    - 实验设计阶段的样本预估偏差
                    """)

        #----------------------------

        # ------- 3.6 分布直方图 -------
        if st.button("刷新页面"):
            st.rerun()
        st.subheader("📉 Distribution Comparison between Groups")
        fig, ax = plt.subplots(figsize=(10, 6))
        ax.hist(val_a, bins=20, alpha=0.5, label=f'{group_a_name} (Group A)', color='r')
        ax.hist(val_b, bins=20, alpha=0.5, label=f'{group_b_name} (Group B)', color='b')
        ax.set_xlabel('Value', fontsize=14)
        ax.set_ylabel('Frequency', fontsize=14)
        ax.set_title('Distribution Comparison', fontsize=18)
        ax.legend()
        st.pyplot(fig)

        # ------- 3.7 多维度下钻分析 -------
        with st.expander("🔍 多维度下钻分析（按类别拆分查看隐藏效应）"):
            st.markdown("选择一列分类变量，系统将自动对每个子群体分别执行AB测试。")
            all_cols = df.columns
            exclude_cols = [group_col, value_col, conv_col]
            drill_cols = [col for col in all_cols if col not in exclude_cols]

            if not drill_cols:
                st.info("当前数据中没有其他分类列可供下钻分析。")
            else:
                selected_drill = st.selectbox("请选择下钻维度（分类列）", options=drill_cols)
                drill_values = df[selected_drill].drop_nulls().unique().to_list()

                if len(drill_values) < 2:
                    st.warning(f"列 '{selected_drill}' 中只有 1 个类别，无法进行下钻对比。")
                else:
                    st.caption(f"共发现 {len(drill_values)} 个子群体，正在分别计算AB效应...")
                    results_data = []

                    for val in drill_values:
                        sub_df = df.filter(pl.col(selected_drill) == val)
                        sub_groups = sub_df[group_col].unique().to_list()
                        if len(sub_groups) < 2:
                            continue
                        sub_a = sub_df.filter(pl.col(group_col) == group_a_name)[value_col].drop_nulls().to_numpy()
                        sub_b = sub_df.filter(pl.col(group_col) == group_b_name)[value_col].drop_nulls().to_numpy()

                        if len(sub_a) < 5 or len(sub_b) < 5:
                            results_data.append({
                                "子群体": str(val),
                                "样本量(A/B)": f"{len(sub_a)}/{len(sub_b)}",
                                "P值": "样本过少",
                                "结论": "⚠️ 不可信"
                            })
                            continue

                        t_stat_sub, p_val_sub = ttest_ind(sub_a, sub_b, equal_var=False)
                        mean_diff = np.mean(sub_b) - np.mean(sub_a)
                        sig = "✅ 显著" if p_val_sub < 0.05 else "❌ 不显著"
                        results_data.append({
                            "子群体": str(val),
                            "样本量(A/B)": f"{len(sub_a)}/{len(sub_b)}",
                            "均值差 (B-A)": f"{mean_diff:.2f}",
                            "P值": f"{p_val_sub:.4f}",
                            "结论": sig
                        })

                    st.session_state['drill_results'] = results_data

                    if results_data:
                        import pandas as pd
                        st.dataframe(pd.DataFrame(results_data))
                        sig_count = sum(1 for r in results_data if "✅" in r["结论"])
                        st.info(f"💡 在 {len(results_data)} 个子群体中，有 {sig_count} 个显示出显著差异（P<0.05）。")
                        if sig_count > 0 and sig_count < len(results_data):
                            st.warning("⚠️ 注意：部分群体显著而部分不显著，可能存在**辛普森悖论**！")
                    else:
                        st.info("没有足够的数据进行下钻分析。")

        # ------- 3.8 实验设计诊断（统计功效）-------
        with st.expander("🔬 实验设计诊断（统计功效与最小样本量）"):
            st.markdown("""
            **统计功效 (Power)**：若真实差异存在，实验能发现它的概率。一般要求 **> 0.8**。  
            **最小样本量**：在给定预期提升下，为了达到目标功效，每组至少需要的样本数。
            """)
            col_p1, col_p2 = st.columns(2)

            with col_p1:
                st.subheader("📊 均值 T 检验 功效")
                if 'cohens_d' in locals() and len(val_a) > 1 and len(val_b) > 1:
                    power_analysis = TTestIndPower()
                    power_mean = power_analysis.power(
                        effect_size=cohens_d,
                        nobs1=len(val_a),
                        alpha=0.05,
                        ratio=len(val_b) / len(val_a)
                    )
                    st.metric("观测功效", f"{power_mean:.3f}")
                    if power_mean < 0.8:
                        st.warning("⚠️ 功效不足 (< 0.8)")
                    else:
                        st.success("✅ 功效充足 (>= 0.8)")
                else:
                    st.info("暂无均值数据用于功效计算。")

            with col_p2:
                st.subheader("📈 比例差异 (卡方) 功效")
                if 'rate_a' in locals() and 'rate_b' in locals() and len(val_a) > 1:
                    try:
                        effect_size_prop = proportion_effectsize(rate_a / 100, rate_b / 100)
                        avg_n = (len(val_a) + len(val_b)) / 2
                        power_prop = NormalIndPower().power(
                            effect_size=effect_size_prop,
                            nobs1=avg_n,
                            alpha=0.05,
                            ratio=1
                        )
                        st.metric("观测功效", f"{power_prop:.3f}")
                        if power_prop < 0.8:
                            st.warning("⚠️ 功效不足 (< 0.8)")
                        else:
                            st.success("✅ 功效充足 (>= 0.8)")
                    except Exception:
                        st.info("比例功效计算失败，可能是转化率极端值导致。")
                else:
                    st.info("暂无比例数据用于功效计算。")

        # ------- 3.9 前瞻性样本量计算 -------
        col_s1, col_s2 = st.columns(2)
        with col_s1:
            default_base = float(rate_a) if 'rate_a' in locals() else 20.0
            target_lift = st.number_input(
                "期望的最小相对提升幅度 (%)",
                min_value=1.0, value=10.0, step=1.0, max_value=100.0,
            ) / 100.0
            base_rate = st.number_input(
                "基准组转化率 (%)",
                min_value=0.1, value=round(default_base, 1), step=0.1, max_value=100.0,
            ) / 100.0

        with col_s2:
            power_target = st.slider(
                "目标统计功效 (Power)",
                min_value=0.5, max_value=0.99, value=0.8, step=0.05
            )
            alpha_target = st.selectbox(
                "显著性水平 (Alpha)",
                options=[0.01, 0.05, 0.10], index=1
            )

        if st.button("🚀 计算所需最小样本量"):
            if base_rate > 0 and target_lift > 0:
                effect_size = proportion_effectsize(
                    base_rate,
                    base_rate * (1 + target_lift),
                )
                if effect_size == 0:
                    st.warning("提升幅度为 0，无法计算样本量。")
                else:
                    sample_size_per_group = NormalIndPower().solve_power(
                        effect_size=effect_size,
                        power=power_target,
                        alpha=alpha_target,
                        ratio=1
                    )
                    required_n = int(np.ceil(sample_size_per_group))
                    total_n = required_n * 2
                    st.success(f"""
                        ✅ 为了有 **{power_target*100:.0f}%** 的把握检测到 **{target_lift*100:.0f}%** 的提升（基准转化率 {base_rate*100:.1f}%）：  
                        - **每组至少需要 {required_n} 个样本**  
                        - **总共需要约 {total_n} 个样本**
                    """)
                    if 'val_a' in locals() and 'val_b' in locals():
                        current_n = min(len(val_a), len(val_b))
                        if current_n < required_n:
                            st.warning(f"📌 当前每组仅有 {current_n} 个样本，建议补充至 {required_n} 个以上。")
                        else:
                            st.info(f"📌 当前每组有 {current_n} 个样本，已满足最低要求。")
            else:
                st.error("请确保基准转化率和提升幅度均大于 0。")

        # ------- 3.10 导出 Word 报告 -------
        with st.expander("📄 导出分析报告（Word 文档）"):
            st.markdown("点击下方按钮，将当前分析结果导出为一份结构化的 Word 报告。")

            def fig_to_bytes(fig):
                buf = BytesIO()
                fig.savefig(buf, format='png', dpi=150, bbox_inches='tight')
                buf.seek(0)
                return buf

            if st.button("📥 生成下载 Word 报告按钮"):
                doc = Document()
                title = doc.add_heading('AB测试分析报告', level=1)
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_paragraph(f'报告生成时间: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
                doc.add_paragraph(f'数据文件: {uploaded_file.name if uploaded_file else "未命名"}')
                doc.add_paragraph("")

                doc.add_heading('一、实验概述', level=2)
                doc.add_paragraph(f'• 分组列：{group_col} ({group_a_name} vs {group_b_name})')
                doc.add_paragraph(f'• 数值指标列：{value_col}')
                doc.add_paragraph(f'• 转化指标列：{conv_col}')
                doc.add_paragraph(f'• 总样本量：{len(val_a) + len(val_b)} (A组{len(val_a)}, B组{len(val_b)})')
                doc.add_paragraph("")

                doc.add_heading('二、T检验结果', level=2)
                doc.add_paragraph(f'• {group_a_name} 平均值：{mean_a:.2f}')
                doc.add_paragraph(f'• {group_b_name} 平均值：{mean_b:.2f}')
                doc.add_paragraph(f'• 均值差（B - A）：{mean_b - mean_a:.2f}')
                doc.add_paragraph(f'• T统计量：{t_stat:.4f}')
                doc.add_paragraph(f'翻译: t 越大，对应的 p 值通常越小，越容易显著')
                doc.add_paragraph(f'• P值：{p_value:.4f}')
                doc.add_paragraph(f'翻译: 两组数值具有{"显著" if p_value < 0.05 else "不显著"}差异')
                doc.add_paragraph(f'• 效应量 (Cohen\'s d)：{cohens_d:.4f}')
                doc.add_paragraph("✅ 结论：两组差异显著" if p_value < 0.05 else "ℹ️ 结论：两组差异不显著")
                doc.add_paragraph("")

                doc.add_heading('三、转化率对比', level=2)
                doc.add_paragraph(f'• {group_a_name} 转化率：{rate_a:.2f}%')
                doc.add_paragraph(f'• {group_b_name} 转化率：{rate_b:.2f}%')
                doc.add_paragraph(f'• 相对提升幅度：{lift:.1f}%' if lift is not None else '• 相对提升幅度：N/A')
                doc.add_paragraph(f'• 卡方检验 P值：{p_value_conv:.4f}' if p_value_conv is not None else '• 卡方检验 P值：未计算')
                doc.add_paragraph(f"翻译:{'显著相关' if p_value_conv < 0.05 else '互相独立'}")
                doc.add_paragraph("")

                doc.add_heading('四、分布对比图', level=2)
                img_bytes = fig_to_bytes(fig)
                doc.add_picture(img_bytes, width=Inches(6))
                doc.add_paragraph("")

                doc.add_heading('五、综合结论', level=2)
                if p_value < 0.05:
                    doc.add_paragraph(f'✅ 数值指标显著优于对照组。')
                else:
                    doc.add_paragraph(f'ℹ️ 数值指标无显著差异。')

                doc.add_heading('附录：分组汇总表', level=2)
                summary_rows = summary.to_pandas().values.tolist()
                table = doc.add_table(rows=1 + len(summary_rows), cols=len(summary.columns))
                table.style = 'Table Grid'
                for j, col_name in enumerate(summary.columns):
                    table.cell(0, j).text = str(col_name)
                for i, row in enumerate(summary_rows):
                    for j, val in enumerate(row):
                        table.cell(i + 1, j).text = str(val)

                # ---- 新增：七、下钻分析结果 ----
                if st.session_state.get('drill_results') and len(st.session_state['drill_results']) > 0:
                    doc.add_heading('七、下钻分析结果', level=2)
                    doc.add_paragraph(f'下钻维度:{st.session_state["drill_results"]}')
                    drill_data = st.session_state['drill_results']
                    #构建表格
                    headers = list(drill_data[0].keys())
                    table = doc.add_table(rows=1+len(drill_data), cols=len(headers))
                    table.style = 'Table Grid'
                    for j, h in enumerate(headers):
                        table.cell(0, j).text = str(h)
                    for i, row in enumerate(drill_data):
                        for j, key in enumerate(headers):
                            table.cell(i+1, j).text = str(row.get(key, ''))
                    doc.add_paragraph("注：'✅ 显著'表示P<0.05；'❌ 不显著'表示P≥0.05；'⚠️ 不可信'表示样本量过少。")
                else:
                    doc.add_heading('七、下钻分析结果', level=2)
                    doc.add_paragraph('未进行下钻分析或无可用的下钻数据')

                doc_bytes = BytesIO()
                doc.save(doc_bytes)
                doc_bytes.seek(0)
                safe_filename = f'AB测试报告_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
                st.download_button(
                    label="📥 点击下载 Word 报告 (.docx)",
                    data=doc_bytes,
                    file_name=safe_filename,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
                st.success("✅ 报告生成完成！点击上方按钮下载。")
    else:
        st.info('请先完成数据列配置')
else:
    st.info('👈 请在左侧上传 CSV 文件')
